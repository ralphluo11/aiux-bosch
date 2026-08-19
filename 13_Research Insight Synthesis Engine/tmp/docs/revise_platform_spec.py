from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SRC = Path("/Users/ZCO9SGH/Desktop/UXGS_Enterprise_Research_Platform_Spec_v0.1.docx")
OUT = Path("output/doc/UXGS_Enterprise_Research_Platform_Spec_v0.2.docx")


def replace_paragraph(doc, old, new):
    for p in doc.paragraphs:
        if p.text.strip() == old:
            p.text = new
            return True
    return False


def set_cell_text(cell, text):
    cell.text = text
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = 0


def insert_decision_block(doc):
    anchor = next(p for p in doc.paragraphs if p.text.strip() == "章节结构")

    heading = doc.add_paragraph("0.1 已确认的范围决策")
    heading.style = "Heading 2"
    anchor._p.addprevious(heading._p)

    intro = doc.add_paragraph(
        "以下决策自 v0.2 起作为集成约束。它们用于消除旧 POC、执行 PRD与平台长期方向之间的歧义；如后续改变，必须通过正式 Scope Gate 更新本 Spec 与对应 PRD。"
    )
    anchor._p.addprevious(intro._p)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["事项", "v0.2 决策", "状态"]
    for i, value in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], value)
    rows = [
        ("当前 Alpha 主线", "已有录音/转写/项目材料 → Evidence → Analysis → Human Review → Structured Delivery", "Approved for planning"),
        ("实时 AI 访谈", "不进入 Alpha 或当前 MVP；MVP 完成后再评估是否立项", "Post-MVP TBC"),
        ("产品 Owner", "本版不指定个人；待治理角色确认后补充", "TBC"),
        ("Owner Portal / Dashboard", "Alpha 仅保留必要的版本与审计记录；完整运营界面进入 MVP 或后续", "Phased"),
        ("Word / PPT", "Alpha 只生成结构化内容与模板字段，不承诺生成正式 .docx/.pptx 文件", "File generation deferred"),
    ]
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            set_cell_text(cells[i], value)
    anchor._p.addprevious(table._tbl)

    status_heading = doc.add_paragraph("0.2 能力状态口径")
    status_heading.style = "Heading 2"
    anchor._p.addprevious(status_heading._p)
    status = doc.add_paragraph(
        "全文涉及能力时使用以下口径：Implemented（代码和测试支持）、Experience Demo（界面或 Mock）、Contract Defined（契约已定义）、Planned for Alpha、Planned for MVP/Pilot、Strategic Direction、TBC。除 Implemented 外均不得对外表述为已完成。"
    )
    anchor._p.addprevious(status._p)


doc = Document(SRC)
OUT.parent.mkdir(parents=True, exist_ok=True)

paragraph_updates = {
    "UXGS 企业级 AI Research 平台": "UXGS Enterprise Intelligence Platform · Research Agent",
    "文档版本：v0.1 — Integration Baseline": "文档版本：v0.2 - Integration Draft",
    "状态：Working Spec；后续与现有 POC、PRD、课件及项目材料合并": "状态：Working Spec - Scope decisions recorded；待与现有母 PRD、执行 PRD及实现状态矩阵完成合并后，方可升级为 Integration Baseline",
    "产品 Owner：Joe / UXGS · GSBDU APAC": "产品 Owner：TBC / UXGS · GSBDU APAC",
    "文档用途：指导 Alpha 建设、团队分工、技术选型、测试验收与后续企业部署；它不是最终 UI 稿，也不替代每个项目自己的 Research Plan。": "文档用途：指导 Alpha 建设、团队分工、候选技术选型、测试验收与后续企业部署；它不是最终 UI 稿，也不替代每个项目自己的 Research Plan。涉及未批准技术、指标或项目事实时，均按 Proposed / TBC 管理。",
    "实时 AI 主持全部访谈；Alpha 先支持已有录音/转写/材料进入后半段研究流程。": "实时 AI 主持访谈不进入 Alpha 或当前 MVP。MVP 完成后基于客户价值、合规、成本和技术可行性另行评估，当前状态为 Post-MVP TBC。",
    "访谈/桌面研究/项目资料的整理、编码、主题分析、洞察生成、机会生成与报告/PPT内容输出。": "已有录音/转写、桌面研究和项目资料的整理、编码、主题分析、洞察生成、机会生成，以及报告/PPT的结构化内容输出。Alpha 不承诺实时 AI 访谈或正式 .docx/.pptx 文件生成。",
    "平台侧 Skills 运行、版本管理、评测、发布与回滚；Owner Portal 与运行 Dashboard。": "平台侧能力分阶段建设：Alpha 只提供最小 Skill Runtime、版本记录和离线评测；完整发布/回滚、Owner Portal 与运行 Dashboard 进入 MVP 或后续阶段。",
    "SharePoint、企业文件源、搜索、报告/PPT导出等连接能力，具体 Connector 取决于 Bosch 批准。": "SharePoint、企业文件源、搜索及正式文件导出属于 Pilot 候选连接能力，具体 Connector、数据边界和上线阶段取决于 Bosch 批准。",
    "Human View：Executive Summary、报告、One-pager、Proposal Deck 内容。": "Human View：Executive Summary、报告、One-pager、Proposal Deck 的结构化内容与模板字段；Alpha 不承诺生成正式 .docx/.pptx 文件。",
    "10.1 推荐 Alpha 技术路线": "10.1 Alpha 候选技术路线（待 Bosch IT / Security / Architecture 评审）",
    "16. Definition of Done 与验收": "16. Definition of Done 与验收（按阶段分层）",
    "下一步不是继续扩架构图": "下一步不是继续扩架构图或恢复实时访谈范围",
}
for old, new in paragraph_updates.items():
    replace_paragraph(doc, old, new)

insert_decision_block(doc)

for table in doc.tables:
    for row in table.rows:
        values = [cell.text.strip() for cell in row.cells]
        if not values:
            continue

        # Benchmark facts not yet supported by an authoritative source in the repository.
        if values[0] == "已有验证线索":
            set_cell_text(row.cells[1], "TBC - 待冰箱项目 Owner 与原始材料确认。候选线索包括实验室场景数、检测窗口及 Precision/Recall 目标；在来源、版本和口径确认前不得作为已批准 Benchmark。")

        # Scope and phase corrections.
        if values[0] == "FR-040":
            set_cell_text(row.cells[2], "Minimal Skill Runtime")
            set_cell_text(row.cells[3], "Alpha 仅按需加载首批已发布/冻结的 Skill，并记录 Skill ID、版本和运行结果；完整 Registry、发布、回滚和退役工作流进入 MVP。客户端不可下载源包。")
        if values[0] == "FR-041":
            set_cell_text(row.cells[1], "Should - MVP")
            set_cell_text(row.cells[3], "MVP 建设最小 Owner Portal，用于测试、审批、发布、回滚和退役 Skill / Knowledge / Template；不作为 Alpha Exit Gate。")
        if values[0] == "FR-060":
            set_cell_text(row.cells[1], "Should - MVP")
            set_cell_text(row.cells[3], "MVP 起显示准确率、引用率、通过率、成本、时延、失败和版本对比；Alpha 允许使用离线评测报告。")
        if values[0] == "FR-070":
            set_cell_text(row.cells[1], "Should - Pilot")
            set_cell_text(row.cells[3], "Pilot 阶段通过 Bosch 批准的 MCP/API 连接文件、SharePoint、搜索与导出工具；Alpha 使用本地/脱敏文件。")

        # Delivery output distinction.
        if values[0] == "Human View":
            set_cell_text(row.cells[1], "Executive Summary、报告、One-pager、Proposal Deck 的结构化内容与模板字段；Alpha 不承诺生成正式 .docx/.pptx 文件。")

        # Architecture table is explicitly proposed.
        if values[0] in {"Frontend", "API", "Agent Runtime", "Structured Data", "Files", "Retrieval", "Queue / Cache", "Observability", "Model Gateway", "CI/CD"}:
            current = row.cells[1].text.strip()
            if not current.startswith("Proposed -"):
                set_cell_text(row.cells[1], f"Proposed - {current}")

        # KPI targets must not look like measured or approved baselines.
        if values[0] in {"结构有效率", "引用可定位率", "无依据事实率", "事实准确率", "专家一次通过率", "关键遗漏率", "假设误报率", "权限违规", "人工节省时间"}:
            current = row.cells[2].text.strip()
            if current and not current.startswith("Proposed target -"):
                set_cell_text(row.cells[2], f"Proposed target - {current}")

        # Roadmap language.
        if values[0].startswith("Phase 1｜Vertical Alpha"):
            set_cell_text(row.cells[2], "材料 → Evidence → Synthesis → Review → One-page Delivery；先支持本地/脱敏数据；采用少量冻结 Skill，不建设完整 Portal。")
        if values[0].startswith("Phase 2｜Governed MVP"):
            set_cell_text(row.cells[2], "SSO/RBAC、项目数据库、持久化状态、最小 Owner Portal、评测和 Trace。实时 AI 访谈不属于本阶段承诺。")
        if values[0].startswith("Phase 4｜Enterprise Hardening"):
            set_cell_text(row.cells[2], "安全/合规、连接器、扩展、DR、SLA、成本模型和运营流程；实时 AI 访谈在 MVP 完成后单独 Go/No-go。")

        # DoD phase labels.
        dod_phase = {
            "DOD-01": "Alpha",
            "DOD-02": "Pilot",
            "DOD-03": "Alpha",
            "DOD-04": "Alpha minimum / MVP full",
            "DOD-05": "MVP/Pilot",
            "DOD-06": "Alpha",
            "DOD-07": "Alpha",
            "DOD-08": "MVP",
            "DOD-09": "Pilot",
            "DOD-10": "All stages",
        }
        if values[0] in dod_phase:
            current = row.cells[1].text.strip()
            set_cell_text(row.cells[1], f"[{dod_phase[values[0]]}] {current}")

        # Immediate actions and dataset numbers are proposed until frozen.
        if values[0] == "5" and len(values) >= 4 and values[1] == "选择首批 6–8 Skills":
            set_cell_text(row.cells[1], "选择首批 4–6 个 Alpha Skills")
            set_cell_text(row.cells[2], "只覆盖材料接入、Evidence、Synthesis、Review 与 Delivery，并为每个 Skill 写测试。")
        if values[0] == "6" and len(values) >= 4 and "基准样本" in values[1]:
            set_cell_text(row.cells[1], "建立候选 30–50 个基准样本")
            set_cell_text(row.cells[2], "先定义样本单位与标注口径；数量为 Proposed Target，经小样本试标后冻结。")

doc.save(OUT)
print(OUT.resolve())
